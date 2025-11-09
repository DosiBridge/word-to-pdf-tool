from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os
import tempfile
import shutil
from pathlib import Path
from typing import Optional
import pypdf
from docx import Document
from pdf2docx import Converter
import fitz  # PyMuPDF
import io

app = FastAPI(title="Document Converter API")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create uploads directory
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


@app.get("/")
async def root():
    return {"message": "Document Converter API"}


@app.post("/api/pdf-to-word")
async def pdf_to_word(file: UploadFile = File(...)):
    """Convert PDF to Word document"""
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")
    
    try:
        # Save uploaded file
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        
        with open(temp_pdf.name, 'wb') as f:
            shutil.copyfileobj(file.file, f)
        
        # Convert PDF to DOCX
        cv = Converter(temp_pdf.name)
        cv.convert(temp_docx.name)
        cv.close()
        
        # Clean up PDF file
        os.unlink(temp_pdf.name)
        
        return FileResponse(
            temp_docx.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=file.filename.replace('.pdf', '.docx')
        )
    except Exception as e:
        if os.path.exists(temp_pdf.name):
            os.unlink(temp_pdf.name)
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


@app.post("/api/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    """Convert Word document to PDF"""
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="File must be a Word document (.docx or .doc)")
    
    try:
        # Save uploaded file
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        with open(temp_docx.name, 'wb') as f:
            shutil.copyfileobj(file.file, f)
        
        # Read Word document and create PDF
        doc = Document(temp_docx.name)
        
        # Create PDF using PyMuPDF
        pdf_doc = fitz.open()
        
        # Extract text from Word document (paragraphs and tables)
        content_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    content_parts.append(row_text)
        
        full_text = "\n".join(content_parts)
        
        # Create PDF with proper pagination
        lines = full_text.split('\n')
        y_position = 50
        page_margin = 50
        line_height = 15
        max_y = 750
        
        page = pdf_doc.new_page()
        
        for line in lines:
            if y_position > max_y:
                page = pdf_doc.new_page()
                y_position = 50
            
            if line.strip():
                page.insert_text(
                    (page_margin, y_position),
                    line[:80],  # Limit line length
                    fontsize=11
                )
                y_position += line_height
        
        # Save PDF
        pdf_doc.save(temp_pdf.name)
        pdf_doc.close()
        
        # Clean up DOCX file
        os.unlink(temp_docx.name)
        
        return FileResponse(
            temp_pdf.name,
            media_type="application/pdf",
            filename=file.filename.replace('.docx', '.pdf').replace('.doc', '.pdf')
        )
    except Exception as e:
        if os.path.exists(temp_docx.name):
            os.unlink(temp_docx.name)
        if os.path.exists(temp_pdf.name):
            os.unlink(temp_pdf.name)
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


@app.post("/api/pdf-to-txt")
async def pdf_to_txt(file: UploadFile = File(...)):
    """Convert PDF to text file"""
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")
    
    try:
        # Save uploaded file
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_txt = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w')
        
        with open(temp_pdf.name, 'wb') as f:
            shutil.copyfileobj(file.file, f)
        
        # Extract text from PDF using PyMuPDF
        pdf_doc = fitz.open(temp_pdf.name)
        text_content = []
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            text_content.append(page.get_text())
        
        pdf_doc.close()
        
        # Write text to file
        full_text = "\n\n".join(text_content)
        temp_txt.write(full_text)
        temp_txt.close()
        
        # Clean up PDF file
        os.unlink(temp_pdf.name)
        
        return FileResponse(
            temp_txt.name,
            media_type="text/plain",
            filename=file.filename.replace('.pdf', '.txt')
        )
    except Exception as e:
        if os.path.exists(temp_pdf.name):
            os.unlink(temp_pdf.name)
        if os.path.exists(temp_txt.name):
            os.unlink(temp_txt.name)
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


@app.post("/api/pdf-unlock")
async def pdf_unlock(
    file: UploadFile = File(...),
    password: Optional[str] = Form(None)
):
    """Remove password protection from PDF"""
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")
    
    try:
        # Save uploaded file
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_unlocked = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        with open(temp_pdf.name, 'wb') as f:
            shutil.copyfileobj(file.file, f)
        
        # Try to unlock PDF
        pdf_reader = pypdf.PdfReader(temp_pdf.name)
        
        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            if password:
                pdf_reader.decrypt(password)
            else:
                # Try common empty password
                try:
                    pdf_reader.decrypt("")
                except:
                    raise HTTPException(
                        status_code=400,
                        detail="PDF is password protected. Please provide the password."
                    )
        
        # Create new PDF without encryption
        pdf_writer = pypdf.PdfWriter()
        
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)
        
        # Don't encrypt - this removes encryption
        
        # Write unlocked PDF
        with open(temp_unlocked.name, 'wb') as f:
            pdf_writer.write(f)
        
        # Clean up original PDF
        os.unlink(temp_pdf.name)
        
        return FileResponse(
            temp_unlocked.name,
            media_type="application/pdf",
            filename=f"unlocked_{file.filename}"
        )
    except pypdf.errors.PdfReadError as e:
        if os.path.exists(temp_pdf.name):
            os.unlink(temp_pdf.name)
        if os.path.exists(temp_unlocked.name):
            os.unlink(temp_unlocked.name)
        raise HTTPException(status_code=400, detail=f"Invalid PDF or incorrect password: {str(e)}")
    except Exception as e:
        if os.path.exists(temp_pdf.name):
            os.unlink(temp_pdf.name)
        if os.path.exists(temp_unlocked.name):
            os.unlink(temp_unlocked.name)
        raise HTTPException(status_code=500, detail=f"Unlocking failed: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

